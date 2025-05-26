# output_writer.py
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

SECTION_HEADERS = {
    "Professional Summary", "Experience", "Education", "Skills",
    "Certifications", "Projects", "Awards", "Publications"
}

def is_section_header(line):
    return line.strip() in SECTION_HEADERS

def write_docx(text: str, filename: str):
    doc = Document()
    lines = text.strip().split("\n")

    for i, line in enumerate(lines):
        line = line.strip()
        if not line:
            continue

        # Header (Name line)
        if i == 0:
            para = doc.add_paragraph(line)
            run = para.runs[0]
            run.bold = True
            run.font.size = Pt(16)
            para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            continue

        # Section Headers
        if is_section_header(line):
            para = doc.add_paragraph()
            run = para.add_run(line)
            run.bold = True
            run.font.size = Pt(12)
            continue

        # Bullet Points
        if line.startswith("- "):
            doc.add_paragraph(line[2:], style='List Bullet')
        else:
            doc.add_paragraph(line)

    doc.save(filename)
def write_txt(text: str, filename: str):
    with open(filename, "w") as f:
        f.write(text)

def write_output(text: str, filename: str, format: str = "docx"):
    if format == "docx":
        write_docx(text, filename)
    elif format == "txt":
        write_txt(text, filename)
    else:
        raise ValueError("Unsupported format")
